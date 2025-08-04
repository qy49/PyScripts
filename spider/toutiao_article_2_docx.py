import os
import random
import re
import shutil
import time

import requests
from DrissionPage import ChromiumPage, ChromiumOptions
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mimetypes
from PIL import Image


def sanitize_filename(filename):
    # 定义非法字符的正则表达式
    illegal_chars = r'[/\\:\*\?"<>\|]'

    # 使用re.sub替换非法字符，这里我们用下划线"_"代替
    clean_filename = re.sub(illegal_chars, '', filename)

    # 去除文件名开头和结尾的点号（.）
    # 去除连续的多个点号
    clean_filename = re.sub(r'^\.', '', clean_filename)
    clean_filename = re.sub(r'\.{2,}', '.', clean_filename)

    # 返回清理后的文件名
    return clean_filename


def get_article_toutiao(url, cookies):
    try:
        response = requests.get(url, cookies=cookies)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.find('h1').text.strip()
        title = sanitize_filename(title)
        cache_path = 'cache'
        if not os.path.exists(cache_path):
            os.mkdir(cache_path)

        for filename in os.listdir(cache_path):
            file_path = os.path.join(cache_path, filename)
            try:
                # 判断是文件还是文件夹，并执行相应的删除操作
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # 删除文件或符号链接
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  # 删除文件夹及其内容
            except Exception as e:
                print(f'删除 {file_path} 时出错: {e}')
        # if not os.path.exists(os.path.join(cache_path, title)):
        #     os.mkdir(os.path.join(cache_path, title))
        # author_tag = soup.find('span', class_='name')
        # author = author_tag.text if author_tag else '未知作者'
        main_tag = soup.find('article')

        tags = main_tag.find_all(recursive=False)
        # 创建doc实例
        doc = Document()
        # 设置标题样式
        title_p = doc.add_heading(text=title, level=1)
        run = title_p.runs[0]
        run.font.name = "Times New Roman"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(20)
        # 设置字体颜色为黑色
        run.font.color.rgb = RGBColor(0, 0, 0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, tag in enumerate(tags):
            img_tag = tag.find('img')
            if img_tag:  # 图片
                try:
                    img_url = img_tag['data-src']
                    response = requests.get(img_url, cookies=cookies)
                    content_type = response.headers['Content-Type']
                    extension = mimetypes.guess_extension(content_type)
                    if content_type == 'image/webp':
                        extension = '.jpg'
                    if extension not in ['.jpg', '.png', '.gif']:
                        continue
                    cache_path = './cache'
                    if not os.path.exists(cache_path):
                        os.mkdir(cache_path)
                    temp_img_name = f'{title}{i+1}{extension}'
                    temp_img_path = os.path.join(cache_path, temp_img_name)
                    with open(temp_img_path, 'wb') as f:
                        f.write(response.content)
                    img = Image.open(temp_img_path)
                    width, height = img.size
                    target_width = 14
                    target_height = target_width * height / width
                    image_para = doc.add_paragraph()
                    run = image_para.add_run()
                    run.add_picture(temp_img_path, width=Cm(target_width), height=Cm(target_height))
                    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    pass
            else:  # 文本
                if not tag.text:
                    continue
                doc.add_paragraph(tag.text)

        target_folder_path = './下载的文章'
        if not os.path.exists(target_folder_path):
            os.mkdir(target_folder_path)
        doc_file_path = os.path.join(target_folder_path, f'{title}.docx')
        doc.save(doc_file_path)
        print(f'下载成功~')
        return 1
    except Exception as e:
        print('下载失败', e)
        return 0


if __name__ == "__main__":
    print('欢迎使用【批量下载头条文章】 ')
    print('作者：B站@自动老李 https://space.bilibili.com/3493107746212758/')
    print('如需定制脚本请加微信：ZiDongLaoLi')

    try:
        co = ChromiumOptions()
        co.set_local_port(9223)
        co.set_argument('--start-maximized')
        page = ChromiumPage(co)
        page.get('https://www.toutiao.com/')

        input('选择分类之后按回车键继续...')

        page.listen.start('https://www.toutiao.com/api/pc/list/feed')

        main_content_node = page.ele('tag=div@class=main-content')
        feed_m_nav_node = main_content_node.ele('tag=div@class=feed-m-nav')
        lis = feed_m_nav_node.eles('tag=li')
        for li in lis:
            if li.attr('aria-pressed') == 'true':
                li.click()
                time.sleep(3)

        index = 1
        while True:
            page.scroll.down(random.randint(400, 600))
            # time.sleep(random.randint(1, 3))
            packet = page.listen.wait(timeout=2)
            if packet:
                try:
                    article_list = packet.response.body['data']
                    for article in article_list:
                        print(f'正在下载第【{index}】篇【{article['title']}】...')
                        get_article_toutiao(f'https://www.toutiao.com/article/{article['item_id']}', page.cookies().as_dict())
                        index += 1
                        time.sleep(random.randint(2, 5))
                except Exception as e:
                    print(e)

    except Exception as e:
        print(e)

    input('按回车键退出...')
